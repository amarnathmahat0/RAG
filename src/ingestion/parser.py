"""Document parser: PDF, DOCX, XLSX, images (OCR), web URLs."""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

from src.utils.logger import get_logger

logger = get_logger(__name__)


@dataclass
class ParsedDocument:
    source: str
    content: str
    metadata: dict[str, Any] = field(default_factory=dict)
    doc_type: str = "text"


class DocumentParser:
    """Parse a file path or URL into a ParsedDocument."""

    def parse(self, source: str) -> ParsedDocument:
        """Route to the correct parser based on extension or URL scheme."""
        if source.startswith("http://") or source.startswith("https://"):
            return self._parse_web(source)
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"Source not found: {source}")
        ext = path.suffix.lower()
        dispatch = {
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".doc": self._parse_docx,
            ".xlsx": self._parse_xlsx,
            ".xls": self._parse_xlsx,
            ".txt": self._parse_text,
            ".md": self._parse_text,
            ".png": self._parse_image,
            ".jpg": self._parse_image,
            ".jpeg": self._parse_image,
            ".tiff": self._parse_image,
            ".bmp": self._parse_image,
        }
        parser_fn = dispatch.get(ext, self._parse_text)
        return parser_fn(source)

    # ── PDF ─────────────────────────────────────────────────────────────────────

    def _parse_pdf(self, path: str) -> ParsedDocument:
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(path)
            blocks: list[str] = []
            metadata: dict[str, Any] = {
                "pages": doc.page_count,
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
            }
            for page_num, page in enumerate(doc, 1):
                text = page.get_text("text")
                if text.strip():
                    blocks.append(f"[Page {page_num}]\n{text.strip()}")
            content = "\n\n".join(blocks)
            logger.info("PDF parsed: %s | pages=%d chars=%d", path, doc.page_count, len(content))
            return ParsedDocument(source=path, content=content, metadata=metadata, doc_type="pdf")
        except ImportError:
            logger.warning("PyMuPDF not installed — falling back to text read.")
            return self._parse_text(path)

    # ── DOCX ────────────────────────────────────────────────────────────────────

    def _parse_docx(self, path: str) -> ParsedDocument:
        try:
            from docx import Document  # python-docx

            doc = Document(path)
            sections: list[str] = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style = para.style.name if para.style else ""
                if style.startswith("Heading"):
                    level = style.split(" ")[-1]
                    prefix = "#" * int(level) if level.isdigit() else "#"
                    sections.append(f"{prefix} {text}")
                else:
                    sections.append(text)
            # Tables
            for table in doc.tables:
                rows: list[list[str]] = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                if rows:
                    header = "| " + " | ".join(rows[0]) + " |"
                    divider = "| " + " | ".join(["---"] * len(rows[0])) + " |"
                    body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
                    sections.append(f"\n{header}\n{divider}\n{body}\n")
            content = "\n\n".join(sections)
            logger.info("DOCX parsed: %s | chars=%d", path, len(content))
            return ParsedDocument(source=path, content=content, metadata={}, doc_type="docx")
        except ImportError:
            logger.warning("python-docx not installed.")
            return self._parse_text(path)

    # ── XLSX ────────────────────────────────────────────────────────────────────

    def _parse_xlsx(self, path: str) -> ParsedDocument:
        try:
            import openpyxl

            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            sections: list[str] = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                sections.append(f"## Sheet: {sheet_name}")
                rows_data: list[list[str]] = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(cells):
                        rows_data.append(cells)
                if rows_data:
                    max_cols = max(len(r) for r in rows_data)
                    for r in rows_data:
                        r.extend([""] * (max_cols - len(r)))
                    header = "| " + " | ".join(rows_data[0]) + " |"
                    divider = "| " + " | ".join(["---"] * max_cols) + " |"
                    body = "\n".join("| " + " | ".join(r) + " |" for r in rows_data[1:])
                    sections.append(f"{header}\n{divider}\n{body}")
            content = "\n\n".join(sections)
            logger.info("XLSX parsed: %s | sheets=%d chars=%d", path, len(wb.sheetnames), len(content))
            return ParsedDocument(source=path, content=content, metadata={"sheets": wb.sheetnames}, doc_type="xlsx")
        except ImportError:
            logger.warning("openpyxl not installed.")
            return ParsedDocument(source=path, content="[XLSX parsing unavailable]", doc_type="xlsx")

    # ── Image / OCR ─────────────────────────────────────────────────────────────

    def _parse_image(self, path: str) -> ParsedDocument:
        try:
            import pytesseract
            from PIL import Image

            img = Image.open(path)
            text = pytesseract.image_to_string(img)
            logger.info("Image OCR: %s | chars=%d", path, len(text))
            return ParsedDocument(
                source=path,
                content=text.strip(),
                metadata={"ocr": True, "size": img.size},
                doc_type="image",
            )
        except ImportError:
            logger.warning("pytesseract/Pillow not available.")
            return ParsedDocument(source=path, content="[OCR unavailable]", doc_type="image")

    # ── Web ─────────────────────────────────────────────────────────────────────

    def _parse_web(self, url: str) -> ParsedDocument:
        try:
            import trafilatura

            downloaded = trafilatura.fetch_url(url)
            if downloaded:
                text = trafilatura.extract(downloaded, include_formatting=True)
                content = text or "[Empty page]"
            else:
                content = "[Failed to fetch page]"
            domain = urlparse(url).netloc
            logger.info("Web parsed: %s | chars=%d", url, len(content))
            return ParsedDocument(
                source=url,
                content=content,
                metadata={"url": url, "domain": domain},
                doc_type="web",
            )
        except ImportError:
            logger.warning("trafilatura not installed.")
            return ParsedDocument(source=url, content="[Web parsing unavailable]", doc_type="web")

    # ── Plain text ───────────────────────────────────────────────────────────────

    def _parse_text(self, path: str) -> ParsedDocument:
        with open(path, encoding="utf-8", errors="replace") as f:
            content = f.read()
        return ParsedDocument(source=path, content=content, metadata={}, doc_type="text")
